# IMPORT LIBRARIES
from openpyxl import load_workbook
from jira import JIRA
from enum import Enum
from datetime import datetime
import os
import xlsxwriter
import docx

# CONNECT TO JIRA AND SET UP DOCUMENTS

# connect to external Jira and sign in as 'user'
print("Connecting and signing user into external Jira...")

# create a Jira object for external jira
# replace 'https://jira.atlassian.com' with the external Jira you need to paste worklogs to
# replace 'username' and 'password' with login details from a user who has the appropriate rights to modify worklogs
jira_external = JIRA('https://jira.atlassian.com',
                     basic_auth=('username', 'password'))
print("Succesfully connected and signed in to external Jira as user.")

# connect to internal Jira and sign in as 'user'.
print("Connecting and signing user into internal Jira...")

# create a Jira object for internal jira
# replace 'https://jira.atlassian.com' with the external Jira you need to pull worklog data from
# replace 'username' and 'password' with login details from a user who has the appropriate rights to modify worklogs
jira_internal = JIRA('https://jira.atlassian.com',
                     basic_auth=('username', 'password'))
print("Succesfully connected and signed in to internal Jira as user.")

# Read excel file "Timesheet_Report.xlsx" with Timesheet data
print("Reading Excel file...")
wb = load_workbook(filename='Timesheet_Report.xlsx', read_only=True)
print("Excel file succesfully loaded.")

# load Excel worksheet 'Timesheet' Excel file stored in 'wb'
print("Loading Excel sheet...")
ws = wb['Timesheet Report']
print("Excel sheet loaded.")

# create word document to write names of summaries that are missing on external JIRA
# this is only needed if there is a mismatch between the names of JIRA issues between internal and external
word_doc_missing_summaries = docx.Document()

# create a second word document to write names of items that have 0 remaining estimate on both Jiras
word_doc_remaining_estimate = docx.Document()

# CONSTANTS
FIRST_ROW = 2
LAST_ROW = ws.max_row+1

COLUMN_KEY_INTERNAL = 1
COLUMN_SUMMARY = 2
COLUMN_ORIGINAL_ESTIMATE = 3
COLUMN_STATUS = 4
COLUMN_DATE = 5
COLUMN_TIMESPENT = 6

HOURS_PER_DAY = 8

# FUNCTIONS


def getCellValue(worksheet,  r, c):
    return worksheet.cell(row=r, column=c).value

# SCRIPT BODY


# search for list of issues assigned to 'user' in the 'Project_Test' project of type = "Story" and store in "list_of_story_issues_JQL"
# list of type <class 'jira.client.ResultList'>
# each item in the list is a type <class 'jira.resources.Issue'>
list_of_story_issues_JQL = jira_external.search_issues("project = Project_Test AND assignee =  user and type = Story", startAt=0, maxResults=10000,
                                                       validate_query=True, fields=None, expand=None, json_result=None)

# print this list of issues of type <class 'jira.resources.Issue'>
print(list_of_story_issues_JQL)

# declare empty list to later store summaries field of JQL (list_of_story_issues_JQL) in
list_of_story_issues_JQL_summary_field = []

# declare empty list to later store summaries of tasks that do not exist on external Jira
list_of_unique_summaries_not_existing_on_external_Jira = []

# loop through every issue in "list_of_story_issues_JQL" list
# print the summary field of that issue (which is of type string)
# add that summary field to the "list_of_story_issues_JQL_summary_field" list
for issue in list_of_story_issues_JQL:
    print(str(issue))
    list_of_story_issues_JQL_summary_field.append(str(issue))

# loop through every row in the excel sheet at column "Summary"
for i in range(FIRST_ROW, LAST_ROW):

    # get the cell value of row 'i',at COLUMN_SUMMARY, split it and get the part at index 0, which is the Jira key of the issue on external Jira
    summary = str(getCellValue(ws, i, COLUMN_SUMMARY)).split(" ")[0]

    # get the cell value of row 'i', at COLUMN_TIMESPENT, and multiply it by 3600, to go from hours to seconds
    #time_spent = str(getCellValue(ws,i,COLUMN_TIMESPENT)) + "h"
    time_spent = float(getCellValue(ws, i, COLUMN_TIMESPENT))*3600

    # get the cell value of row 'i' at COLUMN_DATE, to get the date at which the work was logged
    internaldate = getCellValue(ws, i, COLUMN_DATE)

    # get the cell value of row 'i' at COLUMN_STATUS, to get the status of the parent collection of the task
    status = getCellValue(ws, i, COLUMN_STATUS)

    # get the internal Jira key of the issue and store it in the 'jira_key_internal' variable
    jira_key_internal = str(getCellValue(ws, i, COLUMN_KEY_INTERNAL))

    # search for the issue with 'summary' key on external Jira and include the timetracking fields
    # get the first item of that list, which is the issue (story) on external Jira
    # get the remaining estimate of that issue on external Jira and store it in 'remaining_estimate_external_jira'
    # TRY getting the remaining estimate of that issue on external Jira, because if it does not exist, it will not be possible
    try:
        list_remaining_estimate_external_jira = jira_external.search_issues("issuekey=%s" % (
            summary), startAt=0, maxResults=1, validate_query=True, fields="timetracking", json_result=None)
        list_remaining_estimate_external_jira_first_item = list_remaining_estimate_external_jira[
            0]
        remaining_estimate_external_jira = list_remaining_estimate_external_jira_first_item.fields.timetracking.remainingEstimateSeconds
        print("Remaining estimate on external Jira is %s" %
              (remaining_estimate_external_jira))
    except:
        print("Cannot get remaining estimate from external Jira. The issue likely does not exist yet on external Jira.")

    # use that variable to search for the issue on internal Jira
    # get the first item of that list in order to later get the remaining estimate
    list_issue_internal_jira = jira_internal.search_issues("issuekey=%s" % (
        jira_key_internal), startAt=0, maxResults=1, validate_query=True, fields=["timetracking", "labels"], json_result=None)
    list_issue_internal_jira_first_item = list_issue_internal_jira[0]

    # try getting the remaining estimate from the issue on internal Jira
    # if you cant get the remaining estimate from the issue on internal Jira, it's because it is 'None' and thus == 0
    try:
        remaining_estimate_internal_jira = list_issue_internal_jira_first_item.fields.timetracking.remainingEstimateSeconds
        print("Remaining estimate on internal Jira is %s" %
              (str(remaining_estimate_internal_jira)))
    except:
        print("Remaining estimate on internal Jira is empty")
        remaining_estimate_internal_jira = 0

    # if the summary from that row in the Excel file exists on the external Jira, do this:
    if summary in list_of_story_issues_JQL_summary_field:
        print("Summary exists in external JQL")
        print("Worklog is being added.... ")

        # add a worklog to external Jira issue with 'summary' key
        # timespentSeconds is the time that was spent on the worklog in internal Jira
        # always in timespentSeconds, as easiest conversion
        jira_external.add_worklog(
            summary, timeSpentSeconds=time_spent, started=internaldate)
        print("Worklog added.")

        # if the status of the issue is 'Resolved' on internal Jira according to the data in the Excel, do this:
        if status == "Resolved":
            print("Collection has been resolved.")

            # Transition the issue to 'Done' on external Jira and add a mandatory worklog of 1m spent
            jira_external.transition_issue(summary, "Done", worklog="1m")
            print("Issue transitioned to Done.")

            # note: if the task is marked as Done on external Jira, you can still log time
            # note: if the task is marked as Done on external Jira, you can still transition to 'Done' (the status will stay the same)

            # get the latest worklog of issue with 'summary' key on external Jira and delete it
            # this is the latest worklog you added yourself when transitioning the issue to Done on external Jira
            # set the new remaining estimate to 0m
            last_worklog = jira_external.worklogs(summary)[-1]
            print(last_worklog)
            last_worklog.delete(adjustEstimate="new", newEstimate="0m")

        # if the status of the task on internal Jira is not "Resolved" according to the Excel data, but the remaining estimate of the task on external Jira == 0, do this:
        elif status == "Collection" and remaining_estimate_external_jira == 0:

            # Transition the issue to 'In Progress' on external Jira
            jira_external.transition_issue(summary, "In Progress")
            print("Issue transitioned to In Progress.")

            # check if the remaining estimate on internal JIRA is not == 0
            # if the remaining estimate on internal Jira of the task is larger than 0
            if remaining_estimate_internal_jira > 0:
                # Add that remainign estimate as the remaining estimate on the external task
                list_remaining_estimate_external_jira_first_item.update(update={"timetracking": [
                                                                        {"edit": {"remainingEstimate": remaining_estimate_internal_jira}}]})

            # check if the remaining estimate of the task on internal Jira is equal to zero as well, if so, do this:
            elif remaining_estimate_internal_jira == 0:

                print("%s with internal key %s has 0 remaining estimate in external Jira and internal Jira. Please manually add time." % (
                    str(getCellValue(ws, i, COLUMN_SUMMARY)), str(getCellValue(ws, i, COLUMN_KEY_INTERNAL))))

                # Add a paragraph to the remaining estimate document with a description of the task that needs investigating into the remaining time
                word_doc_remaining_estimate_paragraph = word_doc_remaining_estimate.add_paragraph(
                    str(getCellValue(ws, i, COLUMN_SUMMARY)))
                word_doc_remaining_estimate_paragraph.add_run(
                    "with internal key ")
                word_doc_remaining_estimate_paragraph.add_run(
                    str(getCellValue(ws, i, COLUMN_KEY_INTERNAL)))
                word_doc_remaining_estimate_paragraph.add_run(
                    "has 0 remaining estimate in external Jira and internal Jira. Please manually add time.")

        # check if the status on of the task on internal Jira is still open (for example, if it has re-opened compared to last time), and if so:
        elif status == "Collection":
            # Transition the issue to 'In Progress' on external Jira
            jira_external.transition_issue(summary, "In Progress")
            print("Issue transitioned to In Progress.")

    # if the summary of the task in Excel is not present in the external Jira JQL, do this:
    else:
        print(str(getCellValue(ws, i, COLUMN_SUMMARY)))

        # get the summary of the task, string format it and store it in summary_external_jira
        summary_external_jira = getCellValue(
            ws, i, COLUMN_SUMMARY).replace(" -", ":").replace("_", " ")
        print(summary_external_jira)

        # get the issue internal Jira key from the Excel and store it in:
        jira_key_internal = str(getCellValue(ws, i, COLUMN_KEY_INTERNAL))

        # get the original estimate of the internal Jira issue, convert it into a string and store it in:
        original_estimate_hours_str = str(
            getCellValue(ws, i, COLUMN_ORIGINAL_ESTIMATE))

        # get the original_estimate_hours_str value, remove the 'h', cast it as a float and store it under:
        original_estimate_hours_float = float(
            original_estimate_hours_str.replace("h", ""))

        # get that float number, split it into the amount of hours per working day to get the amount of Mandays and store it in:
        original_estimate_days_float = original_estimate_hours_float/HOURS_PER_DAY

        # convert that number into a string so it can be parsed in the document later
        original_estimate_hours_str = str(original_estimate_days_float)

        # check if the task has not been added to the document yet, by checking if it exists in the list of unique summaries that are not on external Jira.
        # If they are not, do this:
        if summary_external_jira not in list_of_unique_summaries_not_existing_on_external_Jira:
            # add the summary to the list, to make sure it does not get added to the document again.
            list_of_unique_summaries_not_existing_on_external_Jira.append(
                summary_external_jira)

            # add "NotOnExternalJira" label to issue on internal
            list_issue_internal_jira_first_item.fields.labels.append(
                u"NotOnExternalJira")
            list_issue_internal_jira_first_item.update(
                fields={"labels": list_issue_internal_jira_first_item.fields.labels})

            # add the description of the missing task to the document:
            word_doc_missing_summaries_paragraph = word_doc_missing_summaries.add_paragraph(
                summary_external_jira)
            word_doc_missing_summaries_paragraph.add_run(
                " (original estimate: ")
            word_doc_missing_summaries_paragraph.add_run(
                original_estimate_hours_str)
            word_doc_missing_summaries_paragraph.add_run("d,")
            word_doc_missing_summaries_paragraph.add_run(jira_key_internal)
            word_doc_missing_summaries_paragraph.add_run(")")

# Save all the data added to the documents in docx documents
word_doc_missing_summaries.save('MissingTasks.docx')
word_doc_remaining_estimate.save('RemainingTime.docx')
